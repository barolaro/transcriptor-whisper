import io
import json
import os
import tempfile
import uuid
from pathlib import Path

import streamlit as st
from docx import Document
from snowflake.snowpark.context import get_active_session


STAGE_NAME = "TRANSCRIPTOR_AUDIO_STAGE"
SUPPORTED_FORMATS = ["aac", "flac", "m4a", "mp3", "mp4", "ogg", "wav", "webm"]
MAX_FILE_MB = 700

st.set_page_config(page_title="Transcriptor de Audio", page_icon="🎙️")


def create_word_document(text: str, original_name: str) -> bytes:
    document = Document()
    document.add_heading("Transcripción de audio", 0)
    document.add_paragraph(f"Archivo: {original_name}")
    document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def parse_result(raw):
    value = raw
    for _ in range(3):
        if isinstance(value, str):
            value = json.loads(value)
        elif isinstance(value, dict) and "value" in value:
            if value.get("error"):
                raise RuntimeError(value["error"])
            value = value["value"]
        else:
            break
    return value


def format_speaker_transcript(result: dict) -> str:
    segments = result.get("segments") or []
    if not segments:
        return result.get("text", "")
    lines = []
    last_speaker = None
    for segment in segments:
        speaker = segment.get("speaker_label", "HABLANTE")
        text = segment.get("text", "").strip()
        if not text:
            continue
        if speaker == last_speaker and lines:
            lines[-1] += f" {text}"
        else:
            lines.append(f"{speaker}: {text}")
            last_speaker = speaker
    return "\n\n".join(lines)


def transcribe(uploaded_file, identify_speakers: bool):
    session = get_active_session()
    session.sql(
        f"CREATE STAGE IF NOT EXISTS {STAGE_NAME} DIRECTORY=(ENABLE=TRUE)"
    ).collect()

    extension = Path(uploaded_file.name).suffix.lower()
    remote_name = f"audio_{uuid.uuid4().hex}{extension}"
    local_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as tmp:
            tmp.write(uploaded_file.getbuffer())
            local_path = tmp.name

        session.file.put(
            local_path,
            f"@{STAGE_NAME}/{remote_name}",
            auto_compress=False,
            overwrite=True,
        )

        if identify_speakers:
            sql = (
                f"SELECT AI_TRANSCRIBE(TO_FILE('@{STAGE_NAME}', ?), "
                "{'timestamp_granularity': 'speaker'}, TRUE) AS RESULT"
            )
        else:
            sql = (
                f"SELECT AI_TRANSCRIBE(TO_FILE('@{STAGE_NAME}', ?), "
                "OBJECT_CONSTRUCT(), TRUE) AS RESULT"
            )

        raw = session.sql(sql, params=[remote_name]).collect()[0]["RESULT"]
        return parse_result(raw)
    finally:
        try:
            session.sql(f"REMOVE @{STAGE_NAME}/{remote_name}").collect()
        except Exception:
            pass
        if local_path and os.path.exists(local_path):
            os.remove(local_path)


st.title("🎙️ Transcriptor de Audio")
st.write(
    "Transcribe audios directamente con Snowflake AI_TRANSCRIBE, sin descargar "
    "modelos ni consumir la CPU de Streamlit."
)

uploaded_file = st.file_uploader(
    "📂 Selecciona un archivo de audio o video",
    type=SUPPORTED_FORMATS,
)
identify_speakers = st.checkbox(
    "Identificar hablantes",
    value=False,
    help="Disponible para archivos de hasta 60 minutos.",
)

if uploaded_file is not None:
    size_mb = uploaded_file.size / (1024 * 1024)
    st.caption(f"{uploaded_file.name} · {size_mb:.1f} MB")
    if size_mb > MAX_FILE_MB:
        st.error(f"El archivo supera el límite de {MAX_FILE_MB} MB de Snowflake.")
        st.stop()

    if st.button("🚀 Transcribir", type="primary", use_container_width=True):
        try:
            with st.spinner("Subiendo y transcribiendo el archivo…"):
                result = transcribe(uploaded_file, identify_speakers)
            transcript = format_speaker_transcript(result)
            st.session_state["transcript"] = transcript
            st.session_state["audio_duration"] = result.get("audio_duration", 0)
            st.session_state["source_name"] = uploaded_file.name
        except Exception as exc:
            st.error("No fue posible transcribir el archivo.")
            st.exception(exc)

if st.session_state.get("transcript"):
    duration = st.session_state.get("audio_duration", 0)
    st.success(f"Transcripción finalizada · {duration / 60:.1f} minutos")
    edited_text = st.text_area(
        "Puedes corregir el texto antes de descargarlo:",
        value=st.session_state["transcript"],
        height=360,
    )
    source_name = st.session_state.get("source_name", "audio")
    base_name = Path(source_name).stem
    col1, col2 = st.columns(2)
    col1.download_button(
        "📥 Descargar TXT",
        edited_text.encode("utf-8"),
        file_name=f"{base_name}_transcripcion.txt",
        mime="text/plain",
        use_container_width=True,
    )
    col2.download_button(
        "📥 Descargar Word",
        create_word_document(edited_text, source_name),
        file_name=f"{base_name}_transcripcion.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
