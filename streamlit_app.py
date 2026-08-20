import hashlib
import io
import os
import re
import subprocess
import tempfile
from pathlib import Path

import streamlit as st
from docx import Document
from faster_whisper import WhisperModel


st.set_page_config(
    page_title="Transcriptor de Audio",
    page_icon="🎙️",
    layout="centered",
)

SUPPORTED_FORMATS = ["mp3", "wav", "m4a", "mp4", "aac", "ogg", "webm", "flac", "opus"]

# --- Ajustes de robustez para audios largos/pesados ------------------------
# Tope duro de subida (el config.toml manda en la práctica; esto es un aviso).
MAX_UPLOAD_MB = 1000
# Si el audio dura más que esto, se trocea en ventanas para no reventar la RAM
# ni quedar "colgado" en Streamlit Cloud. Por debajo, va en una sola pasada.
SINGLE_PASS_MAX_SECONDS = 5 * 60           # 5 min
# Fragmentos pequeños reducen el uso máximo de RAM en Streamlit gratuito.
CHUNK_TARGET_SECONDS = 4 * 60               # ~4 min por trozo
CHUNK_MAX_SECONDS = 5 * 60                  # nunca más de 5 min sin cortar
SILENCE_NOISE_DB = -30                     # umbral de "silencio"
SILENCE_MIN_DURATION = 0.5                 # duración mínima de silencio (s)

# Perfiles velocidad/calidad.
#   - "base" es el piso razonable: rápido y mucho mejor que "tiny" en español.
#   - "small" es el techo práctico en Streamlit Cloud gratuito (RAM ~1 GB):
#     mucho mejor calidad, un poco más lento.
MODEL_PROFILES = {
    "Rápido": {
        "model": "base",
        "beam_size": 1,
        "best_of": 1,
        "condition_on_previous_text": False,
        "descripcion": "Más veloz. Bien para audio limpio, una sola voz. "
        "Recomendado para audios muy largos.",
    },
    "Recomendado · mejor calidad": {
        "model": "base",
        "beam_size": 3,
        "best_of": 3,
        "condition_on_previous_text": True,
        "descripcion": "Mejor revisión de alternativas usando el mismo modelo liviano. "
        "Más lento en audios largos.",
    },
}

# Pista inicial: orienta al modelo hacia español con puntuación y mayúsculas.
INITIAL_PROMPT = (
    "Transcripción en español de Chile, con puntuación, acentos y "
    "mayúsculas correctas."
)


# ---------------------------------------------------------------------------
# Utilidades de ffmpeg (memoria acotada: todo pasa por disco, no por RAM)
# ---------------------------------------------------------------------------
def _run(cmd: list[str]) -> subprocess.CompletedProcess:
    return subprocess.run(cmd, capture_output=True, text=True)


def save_upload_to_disk(uploaded_file) -> tuple[str, str, float]:
    """Guarda la subida en disco leyéndola por trozos (sin cargar el archivo
    completo en memoria de golpe) y devuelve (ruta, sha256, tamaño_MB)."""
    suffix = Path(uploaded_file.name).suffix.lower() or ".audio"
    hasher = hashlib.sha256()
    size = 0
    uploaded_file.seek(0)
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        while True:
            block = uploaded_file.read(1024 * 1024)  # 1 MB por vez
            if not block:
                break
            hasher.update(block)
            size += len(block)
            tmp.write(block)
        path = tmp.name
    uploaded_file.seek(0)
    return path, hasher.hexdigest(), size / (1024 * 1024)


def to_wav_16k_mono(src_path: str) -> str:
    """Reconvierte a WAV 16 kHz mono (la entrada nativa de Whisper). Reduce
    el uso de memoria y normaliza cualquier formato de entrada."""
    dst = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name
    proc = _run([
        "ffmpeg", "-y", "-i", src_path,
        "-ac", "1", "-ar", "16000", "-c:a", "pcm_s16le", dst,
    ])
    if proc.returncode != 0 or not os.path.exists(dst):
        raise RuntimeError(f"ffmpeg no pudo convertir el audio:\n{proc.stderr[-800:]}")
    return dst


def get_duration(path: str) -> float:
    proc = _run([
        "ffprobe", "-v", "error", "-show_entries", "format=duration",
        "-of", "default=noprint_wrappers=1:nokey=1", path,
    ])
    try:
        return float(proc.stdout.strip())
    except ValueError:
        return 0.0


def detect_silence_midpoints(path: str) -> list[float]:
    """Devuelve los puntos medios de cada silencio: son los mejores lugares
    para cortar sin partir una palabra."""
    proc = _run([
        "ffmpeg", "-i", path,
        "-af", f"silencedetect=noise={SILENCE_NOISE_DB}dB:d={SILENCE_MIN_DURATION}",
        "-f", "null", "-",
    ])
    log = proc.stderr
    starts = [float(x) for x in re.findall(r"silence_start: ([\d.]+)", log)]
    ends = [float(x) for x in re.findall(r"silence_end: ([\d.]+)", log)]
    return sorted((s + e) / 2.0 for s, e in zip(starts, ends))


def plan_chunks(duration: float, silence_mids: list[float]) -> list[tuple[float, float]]:
    """Arma los rangos [(inicio, fin)] cortando en el silencio más cercano al
    objetivo. Si no hay silencio en la ventana, hace un corte duro."""
    ranges: list[tuple[float, float]] = []
    start = 0.0
    while start < duration - 0.05:
        ideal = start + CHUNK_TARGET_SECONDS
        if ideal >= duration:
            ranges.append((start, duration))
            break
        cands = [
            m for m in silence_mids
            if start + 30 < m < start + CHUNK_MAX_SECONDS
        ]
        cut = min(cands, key=lambda m: abs(m - ideal)) if cands else min(ideal, duration)
        ranges.append((start, cut))
        start = cut
    return ranges


def cut_audio_chunk(src_path: str, start: float, end: float) -> str:
    """Convierte solo un fragmento a WAV 16 kHz mono.

    Evita crear una copia WAV completa del audio, que era el principal pico
    de disco y memoria en Streamlit Community Cloud.
    """
    dst = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name
    proc = _run([
        "ffmpeg", "-y", "-ss", f"{start:.3f}", "-to", f"{end:.3f}",
        "-i", src_path, "-vn", "-ac", "1", "-ar", "16000",
        "-c:a", "pcm_s16le", dst,
    ])
    if proc.returncode != 0 or not os.path.exists(dst):
        raise RuntimeError(
            f"ffmpeg no pudo preparar el fragmento:\n{proc.stderr[-500:]}"
        )
    return dst


# ---------------------------------------------------------------------------
# Modelo y transcripción
# ---------------------------------------------------------------------------
@st.cache_resource(show_spinner=False)
def load_model(model_name: str) -> WhisperModel:
    """Carga una sola vez el modelo optimizado para CPU (INT8)."""
    cpu_threads = max(1, min(2, os.cpu_count() or 1))
    return WhisperModel(
        model_name,
        device="cpu",
        compute_type="int8",
        cpu_threads=cpu_threads,
        num_workers=1,
    )


def _transcribe_file(model: WhisperModel, path: str, profile: dict):
    """Generador crudo: entrega (texto_segmento, fin_relativo_al_archivo)."""
    segments, _info = model.transcribe(
        path,
        language="es",
        task="transcribe",
        beam_size=profile["beam_size"],
        best_of=profile["best_of"],
        temperature=[0.0, 0.2, 0.4, 0.6, 0.8, 1.0],
        compression_ratio_threshold=2.4,
        log_prob_threshold=-1.0,
        no_speech_threshold=0.6,
        condition_on_previous_text=profile["condition_on_previous_text"],
        initial_prompt=INITIAL_PROMPT,
        vad_filter=True,
        vad_parameters={"min_silence_duration_ms": 500, "speech_pad_ms": 400},
    )
    for seg in segments:
        yield seg.text.strip(), seg.end


def transcribe_stream(src_path: str, duration: float, profile: dict):
    """Transcribe siempre por fragmentos pequeños y limita el pico de RAM."""
    model = load_model(profile["model"])
    parts: list[str] = []
    total = duration or 0.0

    if total <= 0:
        raise RuntimeError("No fue posible determinar la duración del audio.")

    # Detectar silencios lee el archivo, pero no crea una copia WAV completa.
    mids = detect_silence_midpoints(src_path)
    chunks = plan_chunks(total, mids)

    for chunk_number, (start, end) in enumerate(chunks, start=1):
        chunk_path = None
        try:
            chunk_path = cut_audio_chunk(src_path, start, end)
            for text, seg_end in _transcribe_file(model, chunk_path, profile):
                if text:
                    parts.append(text)
                processed = min(end, start + seg_end)
                progress = min(0.99, processed / total)
                yield (
                    " ".join(parts).strip(),
                    progress,
                    total,
                    chunk_number,
                    len(chunks),
                )
        finally:
            if chunk_path and os.path.exists(chunk_path):
                os.remove(chunk_path)

    yield " ".join(parts).strip(), 1.0, total, len(chunks), len(chunks)


# ---------------------------------------------------------------------------
# Salida Word
# ---------------------------------------------------------------------------
def create_word_document(text: str, original_name: str) -> bytes:
    document = Document()
    document.add_heading("Transcripción de audio", 0)
    document.add_paragraph(f"Archivo: {original_name}")
    document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def reset_result_if_file_changed(file_id: str) -> None:
    if st.session_state.get("current_file_id") != file_id:
        st.session_state.current_file_id = file_id
        for k in ("transcription", "duration", "result_key"):
            st.session_state.pop(k, None)


# ---------------------------------------------------------------------------
# Interfaz
# ---------------------------------------------------------------------------
st.title("🎙️ Transcriptor de Audio")
st.write(
    "Sube un audio y obtén su transcripción en español. "
    "El procesamiento comienza solo cuando presionas **Transcribir**. "
    "Los audios largos se procesan por tramos, así no se cuelga."
)

model_label = st.radio(
    "Velocidad / calidad",
    options=list(MODEL_PROFILES),
    index=1,
    horizontal=True,
)
st.caption(MODEL_PROFILES[model_label]["descripcion"])

uploaded_file = st.file_uploader("📂 Sube tu archivo de audio", type=SUPPORTED_FORMATS)

if uploaded_file is not None:
    # Se guarda a disco por trozos (no carga el archivo completo en RAM).
    src_path, file_id, size_mb = save_upload_to_disk(uploaded_file)
    reset_result_if_file_changed(file_id)
    st.caption(f"{uploaded_file.name} · {size_mb:.1f} MB")

    if size_mb > MAX_UPLOAD_MB:
        st.error(f"El archivo supera el límite de {MAX_UPLOAD_MB} MB.")
        os.remove(src_path)
        st.stop()

    profile = MODEL_PROFILES[model_label]
    result_key = f"{file_id}:{profile['model']}:{profile['beam_size']}"

    if size_mb > 150:
        st.info(
            "Es un audio pesado. En un servidor gratuito puede tardar bastante; "
            "si va muy lento, prueba el modo **Rápido**."
        )

    if st.button("🚀 Transcribir audio", type="primary", use_container_width=True):
        try:
            st.info(
                "La primera vez puede tardar más mientras se descarga y carga "
                "el modelo. Las siguientes son más rápidas."
            )
            bar = st.progress(0.0, text="Preparando el audio…")
            preview = st.empty()

            duration = get_duration(src_path)

            final_text = ""
            for text, progress, total, chunk_no, chunk_count in transcribe_stream(
                src_path, duration, profile
            ):
                final_text = text
                duration = total
                if progress < 1.0:
                    status = (
                        f"Transcribiendo fragmento {chunk_no} de {chunk_count}… "
                        f"{progress * 100:.0f}%"
                    )
                else:
                    status = "Guardando el resultado…"
                bar.progress(progress, text=status)
                if text:
                    preview.markdown(
                        f"""<div style="max-height:200px;overflow-y:auto;
                        padding:0.75rem;border:1px solid rgba(128,128,128,.3);
                        border-radius:.5rem;white-space:pre-wrap;
                        font-size:0.9rem;">{text[-4000:]}</div>""",
                        unsafe_allow_html=True,
                    )

            # Guardar antes de anunciar el 100 % definitivo.
            st.session_state.transcription = final_text
            st.session_state.duration = duration
            st.session_state.result_key = result_key

            preview.empty()
            bar.progress(1.0, text="Transcripción guardada ✅")
        except Exception as exc:
            st.error(
                "No fue posible procesar el audio. Comprueba que el archivo "
                "no esté dañado e inténtalo nuevamente."
            )
            st.exception(exc)
        finally:
            if os.path.exists(src_path):
                os.remove(src_path)

    transcription = st.session_state.get("transcription")
    if transcription and st.session_state.get("result_key") == result_key:
        duration = st.session_state.get("duration", 0)
        st.success(f"✅ Listo. Audio procesado: {duration / 60:.1f} minutos.")
        st.subheader("📄 Texto transcrito")
        edited_text = st.text_area(
            "Puedes corregir el texto antes de descargarlo:",
            value=transcription,
            height=320,
            label_visibility="collapsed",
        )

        base_name = Path(uploaded_file.name).stem
        word_bytes = create_word_document(edited_text, uploaded_file.name)

        col1, col2 = st.columns(2)
        col1.download_button(
            "📥 Descargar TXT",
            data=edited_text.encode("utf-8"),
            file_name=f"{base_name}_transcripcion.txt",
            mime="text/plain",
            use_container_width=True,
        )
        col2.download_button(
            "📥 Descargar Word",
            data=word_bytes,
            file_name=f"{base_name}_transcripcion.docx",
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            use_container_width=True,
        )
    elif (
        st.session_state.get("result_key") == result_key
        and "transcription" in st.session_state
    ):
        st.warning(
            "No se detectó voz. Prueba con el modo Recomendado o revisa el "
            "volumen del audio."
        )
